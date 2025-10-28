"""
CV Parser Module
Extracts and structures text from PDF, DOCX, and TXT files
Includes metadata extraction and section identification
"""

import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
import logging

# Optional imports with graceful fallbacks
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

logger = logging.getLogger(__name__)

class CVParser:
    """
    Multi-format CV parser with metadata extraction and section detection
    """
    
    SECTION_PATTERNS = {
        'contact': r'(?i)(contact|email|phone|address|linkedin|github)',
        'summary': r'(?i)(summary|objective|profile|about)',
        'work_experience': r'(?i)(work\s*experience|employment|professional\s*experience|experience)',
        'projects': r'(?i)(projects|portfolio|work\s*samples)',
        'skills': r'(?i)(skills|technical\s*skills|competencies|technologies)',
        'research': r'(?i)(research|publications|papers|conference)',
        'certifications': r'(?i)(certifications|certificates|training)',
        'awards': r'(?i)(awards|honors|achievements)',
    }
    
    def __init__(self, extract_metadata: bool = True):
        """
        Initialize CV Parser
        
        Args:
            extract_metadata: Whether to extract document metadata
        """
        self.extract_metadata = extract_metadata
        self.supported_formats = ['.pdf', '.docx', '.txt']
        
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse CV from file path
        
        Args:
            file_path: Path to CV file
            
        Returns:
            Parsed CV data with sections and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        extension = path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported format: {extension}")
            
        # Extract raw text based on format
        if extension == '.pdf':
            text, metadata = self._parse_pdf(file_path)
        elif extension == '.docx':
            text, metadata = self._parse_docx(file_path)
        else:  # .txt
            text, metadata = self._parse_txt(file_path)
            
        # Structure text into sections
        sections = self._identify_sections(text)
        
        # Extract dates and timeline
        timeline = self._extract_timeline(text)
        
        # Calculate text statistics
        stats = self._calculate_statistics(text)
        
        return {
            'raw_text': text,
            'sections': sections,
            'metadata': metadata if self.extract_metadata else {},
            'timeline': timeline,
            'statistics': stats,
            'file_info': {
                'name': path.name,
                'size': path.stat().st_size,
                'format': extension,
                'parsed_at': datetime.now().isoformat()
            }
        }
        
    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text and metadata from PDF
        
        Uses both PyPDF2 and pdfplumber for robust extraction
        """
        text_pypdf = []
        text_plumber = []
        metadata = {}
        
        # Try PyPDF2 first
        if HAS_PYPDF2:
            try:
                with open(file_path, 'rb') as file:
                    pdf = PyPDF2.PdfReader(file)
                    
                    # Extract metadata
                    if self.extract_metadata and pdf.metadata:
                        metadata = {
                            'author': pdf.metadata.get('/Author', ''),
                            'creation_date': str(pdf.metadata.get('/CreationDate', '')),
                            'modification_date': str(pdf.metadata.get('/ModDate', '')),
                            'producer': pdf.metadata.get('/Producer', ''),
                            'title': pdf.metadata.get('/Title', ''),
                            'subject': pdf.metadata.get('/Subject', ''),
                        }
                    
                    # Extract text from all pages
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_pypdf.append(page_text)
                        except Exception as e:
                            logger.warning(f"PyPDF2 failed on page {page_num}: {e}")
                            
            except Exception as e:
                logger.error(f"PyPDF2 extraction failed: {e}")
        else:
            logger.warning("PyPDF2 not available")
            
        # Try pdfplumber for better layout preservation
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_plumber.append(page_text)
                        except Exception as e:
                            logger.warning(f"pdfplumber failed on page: {e}")
                            
            except Exception as e:
                logger.error(f"pdfplumber extraction failed: {e}")
        else:
            logger.warning("pdfplumber not available")
            
        # Combine and choose best extraction
        text = '\n'.join(text_plumber) if text_plumber else '\n'.join(text_pypdf)
        
        if not text:
            # Fallback: try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                    logger.warning("Used fallback text extraction for PDF")
            except:
                raise ValueError("Could not extract text from PDF - install PyPDF2 or pdfplumber")
            
        return text, metadata
        
    def _parse_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text and metadata from DOCX
        """
        if not HAS_DOCX:
            # Fallback: try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                    logger.warning("python-docx not available, used fallback text extraction")
                    return text, {}
            except:
                raise ValueError("Could not extract text from DOCX - install python-docx")
        
        doc = Document(file_path)
        text = []
        metadata = {}
        
        # Extract metadata
        if self.extract_metadata:
            core_props = doc.core_properties
            metadata = {
                'author': core_props.author or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'title': core_props.title or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'last_modified_by': core_props.last_modified_by or ''
            }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
                    
        return '\n'.join(text), metadata
        
    def _parse_txt(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from TXT file
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
            
        metadata = {
            'encoding': 'utf-8',
            'file_size': Path(file_path).stat().st_size
        }
        
        return text, metadata
        
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify and extract CV sections using regex patterns
        """
        sections = {}
        lines = text.split('\n')
        
        current_section = None
        section_content = []
        
        for line in lines:
            # Check if line matches any section header
            section_found = False
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line.strip()) and len(line.strip()) < 50:
                    # Save previous section
                    if current_section and section_content:
                        sections[current_section] = '\n'.join(section_content)
                    
                    # Start new section
                    current_section = section_name
                    section_content = []
                    section_found = True
                    break
                    
            if not section_found and current_section:
                section_content.append(line)
                
        # Save last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
            
        # If no sections identified, treat as single block
        if not sections:
            sections['full_text'] = text
            
        return sections
        
    def _extract_timeline(self, text: str) -> List[Dict]:
        """
        Extract dates and create timeline
        """
        timeline = []
        
        # Date patterns
        date_patterns = [
            r'(\d{4})\s*[-–—]\s*(\d{4}|\bpresent\b|\bcurrent\b)',
            r'(\w{3,9})\s*(\d{4})\s*[-–—]\s*(\w{3,9})\s*(\d{4})',
            r'(\d{1,2})[/-](\d{4})\s*[-–—]\s*(\d{1,2})[/-](\d{4})',
            r'(\d{4})[/-](\d{1,2})\s*[-–—]\s*(\d{4})[/-](\d{1,2})'
        ]
        
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Get surrounding context
                start = max(0, match.start() - 100)
                end = min(len(text), match.end() + 100)
                context = text[start:end].replace('\n', ' ')
                
                timeline.append({
                    'date_string': match.group(),
                    'context': context,
                    'position': match.start()
                })
                
        # Sort by position in document
        timeline.sort(key=lambda x: x['position'])
        
        return timeline
        
    def _calculate_statistics(self, text: str) -> Dict:
        """
        Calculate text statistics for credibility analysis
        """
        words = text.split()
        lines = text.split('\n')
        
        # Buzzword detection
        buzzwords = [
            'synergy', 'leverage', 'innovative', 'cutting-edge', 'revolutionary',
            'disruptive', 'transformative', 'passionate', 'driven', 'results-oriented',
            'thought leader', 'guru', 'ninja', 'rockstar', 'unicorn'
        ]
        
        buzzword_count = sum(1 for word in words if word.lower() in buzzwords)
        
        # URL detection
        url_pattern = r'https?://[^\s]+'
        urls = re.findall(url_pattern, text)
        
        # Email detection (for artifact verification)
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        emails = re.findall(email_pattern, text)
        
        # Number/metric detection
        metric_pattern = r'\d+[%$]?|\d+\.\d+[%$]?'
        metrics = re.findall(metric_pattern, text)
        
        return {
            'word_count': len(words),
            'line_count': len(lines),
            'char_count': len(text),
            'buzzword_count': buzzword_count,
            'buzzword_density': buzzword_count / len(words) if words else 0,
            'url_count': len(urls),
            'urls': urls,
            'email_count': len(emails),
            'metric_count': len(metrics),
            'metrics_found': list(set(metrics))[:20],  # Top 20 unique metrics
            'avg_word_length': sum(len(word) for word in words) / len(words) if words else 0
        }
        
    def validate_structure(self, parsed_cv: Dict) -> Dict[str, Any]:
        """
        Validate CV structure and completeness
        
        Returns validation report
        """
        issues = []
        warnings = []
        
        # Check for required sections
        required_sections = ['work_experience', 'skills']
        for section in required_sections:
            if section not in parsed_cv['sections'] or not parsed_cv['sections'][section].strip():
                issues.append(f"Missing or empty section: {section}")
                
        # Check for suspicious patterns
        if parsed_cv['statistics']['buzzword_density'] > 0.05:
            warnings.append(f"High buzzword density: {parsed_cv['statistics']['buzzword_density']:.2%}")
            
        # Check for timeline gaps
        timeline = parsed_cv['timeline']
        if len(timeline) > 1:
            # Simple gap detection (can be enhanced)
            date_strings = [t['date_string'] for t in timeline]
            if any('present' in ds.lower() or 'current' in ds.lower() for ds in date_strings):
                pass  # Currently employed
            else:
                warnings.append("No current position indicated")
                
        return {
            'valid': len(issues) == 0,
            'issues': issues,
            'warnings': warnings,
            'completeness_score': self._calculate_completeness(parsed_cv)
        }
        
    def _calculate_completeness(self, parsed_cv: Dict) -> float:
        """
        Calculate CV completeness score
        """
        score = 0.0
        max_score = 100.0
        
        # Section presence (60 points)
        important_sections = {
            'work_experience': 25,
            'skills': 20,
            'projects': 15
        }
        
        for section, points in important_sections.items():
            if section in parsed_cv['sections'] and parsed_cv['sections'][section]:
                score += points
                
        # Metrics presence (20 points)
        if parsed_cv['statistics']['metric_count'] > 0:
            score += min(20, parsed_cv['statistics']['metric_count'] * 2)
            
        # Links/artifacts (20 points)  
        if parsed_cv['statistics']['url_count'] > 0:
            score += min(20, parsed_cv['statistics']['url_count'] * 5)
            
        return min(100, (score / max_score) * 100)